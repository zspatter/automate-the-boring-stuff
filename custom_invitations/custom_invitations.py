#! /usr/bin/env python3
# custom_invitations.py -

from os.path import abspath, join

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def generate_invitations(textfile, filename):
    """
    Generates a custom invitation for each name in the given text file.
    Each individual invitation corresponds to a single page in the .docx

    :param str textfile: path to .txt file
    :param str filename: path to .docx file
    """
    doc = docx.Document()

    with open(textfile, 'r') as guests:
        # generates a custom invitation for each guest
        for guest in guests:
            style_run(run=center_paragraph(document=doc).add_run(
                'It would be a pleasure to have the company of'),
                is_bold=True,
                is_italic=True,
                font_size=13)
            style_run(run=center_paragraph(document=doc).add_run(guest.strip()),
                      is_bold=True,
                      font_size=15)
            style_run(run=center_paragraph(document=doc).add_run(
                'at 11101 Memory lane on the evening of'),
                is_bold=True,
                is_italic=True)
            style_run(run=center_paragraph(document=doc).add_run('April 31st'))
            style_run(run=center_paragraph(document=doc).add_run("at 24 O'Clock"),
                      is_bold=True,
                      is_italic=True)

            doc.add_page_break()

    doc.save(filename)


def center_paragraph(document):
    """
    Adds new paragraph and sets alignment to be centered

    :param Document document: docx document
    :return: centered paragraph (no text)
    """
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return paragraph


def style_run(run, is_bold=None, is_italic=None, font_size=12):
    """
    Applies rich text styles to the given Run text object.

    :param Run run: run of text to be styled
    :param bool is_bold: indicates if run should be bold
    :param bool is_italic: indicates if run should be italic
    :param int font_size: indicates font size
    """
    run.font.bold = is_bold
    run.font.italic = is_italic
    run.font.size = Pt(font_size)


if __name__ == '__main__':
    path = join(abspath('.'), 'sample_files')
    generate_invitations(textfile=join(path, 'guests.txt'),
                         filename=join(path, 'invitations.docx'))
